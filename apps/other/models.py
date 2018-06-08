from django.db import models
from django.utils.translation import ugettext_lazy as _

from django.db.models.signals import post_save
from django.dispatch import receiver

from django.conf import settings
import time

from docx import Document


def generate_new_filename(instance, filename):
    fullpath = f"other/{time.strftime('%Y-%m')}/{filename}"
    return fullpath


class PPS(models.Model):
    CONDITIONS = (
        ("0", "штатный"),
        ("1", "внутренний совместитель"),
        ("2", "внешний совместитель"),
        ("3", "по договору")
    )
    name = models.TextField(_("ФИО"), max_length=1024, null=True, blank=True)
    conditions = models.TextField(_("Условия привлечения"), max_length=1024, null=True, blank=True)
    position = models.TextField(_("Должность, ученая степень, ученое звание"), max_length=1024, null=True, blank=True)
    disciplines = models.TextField(_("Перечень читаемых дисциплин"), max_length=1024, null=True, blank=True)
    education = models.TextField(_(
        "Уровень образования, наименование специальности, направления подготовки, наименование присвоенной квалификации "),
        max_length=1024, null=True, blank=True)
    ppk = models.TextField(_("Сведения о дополнительном профессиональном образовании"), max_length=1024, null=True,
                           blank=True)
    weight = models.TextField(
        _("Объем учебной нагрузки по дисциплине, практикам, государственной итоговой аттестации (доля ставки)"),
        max_length=1024, null=True, blank=True)
    experience = models.TextField(_(
        "Стаж практической работы по профилю образовательной программы в профильных организациях с указанием периода работы и должности"),
        max_length=1024, null=True, blank=True)

    created_at = models.DateTimeField(auto_now_add=True, blank=True, null=True)
    updated_at = models.DateTimeField(auto_now=True, blank=True, null=True)

    def __str__(self):
        return u"{}".format(self.name)

    class Meta:
        verbose_name = 'преподаватель'
        verbose_name_plural = 'преподаватели'


class File(models.Model):
    file = models.FileField("Таблица преподавателей", upload_to=generate_new_filename)
    done = models.BooleanField("Обработано", default=False)
    created_at = models.DateTimeField(auto_now_add=True, blank=True, null=True)
    updated_at = models.DateTimeField(auto_now=True, blank=True, null=True)


@receiver(post_save, sender=File, dispatch_uid="add_pps")
def add_pps(sender, instance, **kwargs):
    document = Document(f"{settings.MEDIA_ROOT}/{instance.file}")
    tables = document.tables
    for table in tables:
        for ri, row in enumerate(table.rows):
            if ri != 0 and len("".join([p.text for p in row.cells[1].paragraphs])) > 2 and len(row.cells) == 9:
                pps = PPS(
                    name="\n\n".join([p.text for p in row.cells[1].paragraphs]),
                    conditions="\n\n".join([p.text for p in row.cells[2].paragraphs]),
                    position="\n\n".join([p.text for p in row.cells[3].paragraphs]),
                    disciplines="\n\n".join([p.text for p in row.cells[4].paragraphs]),
                    education="\n\n".join([p.text for p in row.cells[5].paragraphs]),
                    ppk="\n\n".join([p.text for p in row.cells[6].paragraphs]),
                    weight="\n\n".join([p.text for p in row.cells[7].paragraphs]),
                    experience="\n\n".join([p.text for p in row.cells[8].paragraphs])
                )
                pps.save()
